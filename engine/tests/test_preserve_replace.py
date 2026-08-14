"""Tests for format-preserving content replacement engine."""

import io
import os
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from app.core.preserve_replace import (
    replace_preserving_format, PLACEHOLDER_PATTERN, ReplacementError,
)
from app.core.preserve_replace_images import (
    replace_images_preserving_geometry, replace_embedded_objects,
)
from app.core.extractor import extract_format


def _make_temp_docx(name="tpl.docx") -> str:
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    return tmp.name


# ─── 4.1 Text replacement ────────────────────────────────────

class TestTextReplacement:
    """Text replacement preserves run/paragraph formatting."""

    def setup_method(self):
        self.tpl = _make_temp_docx()
        doc = DocxDocument()
        p = doc.add_paragraph()
        run = p.add_run("客户：")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        run2 = p.add_run("{{customer_name}}")
        run2.bold = True
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(self.tpl)
        self.out = _make_temp_docx("out.docx")

    def teardown_method(self):
        for p in (self.tpl, self.out):
            Path(p).unlink(missing_ok=True)

    def test_single_run_replacement_preserves_format(self):
        replace_preserving_format(self.tpl, self.out, {"customer_name": "张权"})
        doc = DocxDocument(self.out)
        text = doc.paragraphs[0].text
        assert "客户：张权" == text

        # Verify formatting preserved
        run = doc.paragraphs[0].runs[1]
        assert run.bold is True
        assert run.font.size.pt == 14.0
        assert run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_placeholder_pattern(self):
        m = PLACEHOLDER_PATTERN.search("{{customer_name}}")
        assert m and m.group(1) == "customer_name"


# ─── 4.2 Image replacement ───────────────────────────────────

class TestImageReplacement:
    """Image replacement preserves geometry (extent/anchor)."""

    def setup_method(self):
        self.tpl = _make_temp_docx()
        doc = DocxDocument()
        img = io.BytesIO()
        Image.new("RGB", (100, 40), (0, 100, 200)).save(img, "PNG")
        img.seek(0)
        run = doc.add_paragraph().add_run()
        # Set alt text as placeholder key
        from docx.oxml.ns import qn as _qn
        run.add_picture(img, width=Cm(5))
        # Set docPr name to the placeholder key
        for drawing in run._r.iter():
            if drawing.tag.endswith('}docPr'):
                drawing.set('name', 'logo')
                break
        doc.save(self.tpl)
        self.out = _make_temp_docx("out.docx")

        # Read original geometry
        self.orig_doc = DocxDocument(self.tpl)

    def teardown_method(self):
        for p in (self.tpl, self.out):
            Path(p).unlink(missing_ok=True)

    def test_image_replacement_preserves_extent(self):
        # Capture original extent
        orig_extents = []
        orig_doc = DocxDocument(self.tpl)
        for blip in orig_doc.part.element.iter():
            if blip.tag.endswith('}extent'):
                orig_extents.append((blip.get('cx'), blip.get('cy')))
        assert orig_extents, "should have extent elements"

        # New image (different content, same size intent)
        new_img = io.BytesIO()
        Image.new("RGB", (100, 40), (0, 200, 100)).save(new_img, "PNG")

        replace_images_preserving_geometry(self.tpl, self.out, {"logo": new_img.getvalue()})

        # Verify extent preserved
        out_doc = DocxDocument(self.out)
        out_extents = []
        for blip in out_doc.part.element.iter():
            if blip.tag.endswith('}extent'):
                out_extents.append((blip.get('cx'), blip.get('cy')))
        assert out_extents == orig_extents


# ─── 4.3 Embedded object ─────────────────────────────────────

class TestEmbeddedObject:
    """Embedded object replacement preserves container."""

    def test_no_objects_raises_error(self):
        self.tpl = _make_temp_docx()
        doc = DocxDocument()
        doc.add_paragraph("No OLE here")
        doc.save(self.tpl)
        self.out = _make_temp_docx("out.docx")
        with pytest.raises(Exception):
            replace_embedded_objects(self.tpl, self.out, {"chart": b"data"})
        Path(self.tpl).unlink(missing_ok=True)
        Path(self.out).unlink(missing_ok=True)


# ─── 4.4 Table replacement ───────────────────────────────────

class TestTableReplacement:
    """Table replacement preserves table style and column widths."""

    def setup_method(self):
        self.tpl = _make_temp_docx()
        doc = DocxDocument()
        table = doc.add_table(rows=3, cols=3, style="Table Grid")
        headers = ["项目", "数量", "金额"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        table.cell(1, 0).text = "{{item}}"
        table.cell(1, 1).text = "{{qty}}"
        table.cell(1, 2).text = "{{amount}}"
        table.cell(2, 0).text = "模板行2"
        doc.save(self.tpl)
        self.out = _make_temp_docx("out.docx")

    def teardown_method(self):
        for p in (self.tpl, self.out):
            Path(p).unlink(missing_ok=True)

    def test_table_style_preserved(self):
        # Capture table style name and grid before
        before = DocxDocument(self.tpl)
        b_table = before.tables[0]
        style_before = b_table.style.name if b_table.style else None
        grid_before = len(b_table._tbl.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblGrid'))

        replace_preserving_format(self.tpl, self.out, {
            "item": "服务费", "qty": "1", "amount": "10000",
        })

        after = DocxDocument(self.out)
        a_table = after.tables[0]
        style_after = a_table.style.name if a_table.style else None
        grid_after = len(a_table._tbl.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblGrid'))

        assert style_before == style_after
        assert grid_before == grid_after
        # Verify cell content replaced
        assert "服务费" in a_table.cell(1, 0).text
        assert "模板行2" in a_table.cell(2, 0).text


# ─── 4.5 End-to-end ──────────────────────────────────────────

class TestEndToEnd:
    """Full pipeline: template -> replacement -> format fingerprint unchanged."""

    def setup_method(self):
        self.tpl = _make_temp_docx()
        doc = DocxDocument()
        doc.add_paragraph("合同编号：{{contract_id}}")
        doc.add_paragraph("甲方：{{customer_name}}")
        table = doc.add_table(rows=2, cols=2, style="Table Grid")
        table.cell(0, 0).text = "字段"
        table.cell(0, 1).text = "值"
        table.cell(1, 0).text = "{{key}}"
        table.cell(1, 1).text = "{{value}}"
        doc.save(self.tpl)
        self.out = _make_temp_docx("out.docx")

        # Compute format fingerprint of template
        self.spec = extract_format(self.tpl)

    def teardown_method(self):
        for p in (self.tpl, self.out):
            Path(p).unlink(missing_ok=True)

    def test_format_fingerprint_unchanged(self):
        replace_preserving_format(self.tpl, self.out, {
            "contract_id": "HT-001",
            "customer_name": "测试公司",
            "key": "金额",
            "value": "100000",
        })

        out_spec = extract_format(self.out)
        # Page setup and style fingerprint should be identical
        assert out_spec.fingerprint() == self.spec.fingerprint()
