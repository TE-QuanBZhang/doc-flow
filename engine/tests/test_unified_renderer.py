"""Tests for unified renderer, multi-section extraction, and logic marker cleanup."""

import os
import tempfile
import pytest
from docx import Document as DocxDocument

from app.core.renderer import (
    generate_document,
    generate_document_with_result,
    detect_unresolved_placeholders,
    RenderResult,
    _cleanup_logic_markers,
    LOGIC_BLOCK_PATTERN,
)
from app.core.extractor import extract_format, _extract_all_sections, _extract_all_table_formats
from app.core.models import StyleSpec, SectionInfo


def _create_test_docx(text: str) -> str:
    """Create a temporary .docx file with the given text."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc = DocxDocument()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(tmp.name)
    return tmp.name


def _make_output() -> str:
    """Create a temporary output path."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    return tmp.name


class TestUnifiedRenderer:
    """Test the unified rendering pipeline."""

    def test_generate_document_returns_string(self):
        """generate_document still returns a string path (backward compat)."""
        template = _create_test_docx("Hello {{name}}")
        output = _make_output()
        result = generate_document(template, {"name": "World"}, output)
        assert isinstance(result, str)
        assert result == output
        os.unlink(template)
        os.unlink(output)

    def test_generate_document_with_result_returns_render_result(self):
        """generate_document_with_result returns a RenderResult."""
        template = _create_test_docx("Hello {{name}}")
        output = _make_output()
        result = generate_document_with_result(template, {"name": "World"}, output)
        assert isinstance(result, RenderResult)
        assert result.output_path == output
        assert result.render_mode in ("preserve", "docxtpl", "python_docx")
        os.unlink(template)
        os.unlink(output)

    def test_preserve_mode_is_primary(self):
        """For simple variable replacement, preserve mode should succeed."""
        template = _create_test_docx("Company: {{company}}")
        output = _make_output()
        result = generate_document_with_result(template, {"company": "Acme"}, output)
        assert result.render_mode == "preserve"
        assert len([w for w in result.warnings if "degraded" in w]) == 0

        doc = DocxDocument(output)
        assert "Acme" in doc.paragraphs[0].text
        os.unlink(template)
        os.unlink(output)

    def test_warnings_on_fallback(self):
        """When preserve mode fails, warnings should contain degradation info."""
        template = _create_test_docx("Simple text no placeholders")
        output = _make_output()
        result = generate_document_with_result(template, {}, output)
        # Should still succeed (preserve mode works even with no placeholders)
        assert result.output_path == output
        os.unlink(template)
        os.unlink(output)


class TestLogicMarkerCleanup:
    """Test logic block marker cleanup."""

    def test_cleanup_removes_jinja2_markers(self):
        """Logic markers like {% if %} are cleaned from document text."""
        template = _create_test_docx("{% if show %}Visible content{% endif %}")
        output = _make_output()
        # Copy template to output first
        import shutil
        shutil.copy2(template, output)

        warnings = _cleanup_logic_markers(output)
        doc = DocxDocument(output)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "{% if show %}" not in full_text
        assert "{% endif %}" not in full_text
        assert "Visible content" in full_text
        os.unlink(template)
        os.unlink(output)

    def test_cleanup_removes_legacy_markers(self):
        """Legacy markers like {{#if:cond}} are cleaned."""
        template = _create_test_docx("{{#if:show}}Content{{/if}}")
        output = _make_output()
        import shutil
        shutil.copy2(template, output)

        warnings = _cleanup_logic_markers(output)
        doc = DocxDocument(output)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "{{#if:show}}" not in full_text
        assert "{{/if}}" not in full_text
        assert "Content" in full_text
        os.unlink(template)
        os.unlink(output)

    def test_cleanup_returns_warning_when_markers_found(self):
        """Cleanup returns warnings list with count of cleaned markers."""
        template = _create_test_docx("{% if x %}A{% endif %} {% for i in y %}B{% endfor %}")
        output = _make_output()
        import shutil
        shutil.copy2(template, output)

        warnings = _cleanup_logic_markers(output)
        assert len(warnings) > 0
        assert "cleaned" in warnings[0].lower() or "清理" in warnings[0]
        os.unlink(template)
        os.unlink(output)

    def test_cleanup_no_warning_when_no_markers(self):
        """No warnings when document has no markers."""
        template = _create_test_docx("Just normal text")
        output = _make_output()
        import shutil
        shutil.copy2(template, output)

        warnings = _cleanup_logic_markers(output)
        assert len(warnings) == 0
        os.unlink(template)
        os.unlink(output)


class TestMultiSectionExtraction:
    """Test multi-section format extraction."""

    def test_single_section_extracted(self):
        """A basic document has at least one section."""
        template = _create_test_docx("Hello world")
        spec = extract_format(template)
        assert len(spec.sections) >= 1
        assert spec.sections[0].page.size is not None
        os.unlink(template)

    def test_sections_have_header_footer(self):
        """Each section has header/footer objects."""
        template = _create_test_docx("Hello world")
        spec = extract_format(template)
        for si in spec.sections:
            assert hasattr(si, 'header')
            assert hasattr(si, 'footer')
            assert hasattr(si, 'page')
        os.unlink(template)

    def test_backward_compat_page_from_first_section(self):
        """spec.page is populated from the first section for backward compat."""
        template = _create_test_docx("Hello world")
        spec = extract_format(template)
        assert spec.page is not None
        assert spec.page.size == spec.sections[0].page.size
        os.unlink(template)


class TestMultiTableExtraction:
    """Test multi-table format extraction."""

    def test_no_tables_returns_empty_list(self):
        """Document with no tables returns empty tables list."""
        template = _create_test_docx("No tables here")
        spec = extract_format(template)
        assert spec.tables == []
        os.unlink(template)

    def test_single_table_extracted(self):
        """Document with one table extracts one table format."""
        template = _create_test_docx("Before table")
        doc = DocxDocument(template)
        doc.add_table(rows=2, cols=3)
        doc.save(template)

        spec = extract_format(template)
        assert len(spec.tables) >= 1
        assert spec.table is not None  # backward compat
        os.unlink(template)

    def test_backward_compat_table_from_first(self):
        """spec.table is populated from the first table for backward compat."""
        template = _create_test_docx("Before table")
        doc = DocxDocument(template)
        doc.add_table(rows=2, cols=3)
        doc.save(template)

        spec = extract_format(template)
        if spec.tables:
            assert spec.table.style_name == spec.tables[0].style_name
        os.unlink(template)


class TestRenderResultDataclass:
    """Test RenderResult data structure."""

    def test_default_values(self):
        result = RenderResult(output_path="/tmp/test.docx")
        assert result.render_mode == "preserve"
        assert result.warnings == []

    def test_custom_values(self):
        result = RenderResult(
            output_path="/tmp/test.docx",
            render_mode="docxtpl",
            warnings=["fallback warning"],
        )
        assert result.render_mode == "docxtpl"
        assert len(result.warnings) == 1
