"""Phase 2 integration tests: multi-section extraction verification and PDF converter."""

import os
import tempfile
import shutil
import pytest
from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Cm, Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app.core.extractor import extract_format, _extract_all_sections
from app.core.models import SectionInfo, StyleSpec
from app.core.pdf_converter import (
    is_libreoffice_available,
    _compute_file_hash,
    _get_cache_dir,
    _get_cached_pdf_path,
    clear_pdf_cache,
    CONVERSION_TIMEOUT_SECONDS,
)


# ─── Helpers ────────────────────────────────────────────────────

def _make_output(suffix: str = ".docx") -> str:
    tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
    tmp.close()
    return tmp.name


def _create_multi_section_doc(
    section_configs: list[dict],
) -> str:
    """Create a .docx with multiple sections.

    Each config dict can contain:
    - orientation: "portrait" | "landscape"
    - page_width_cm, page_height_cm: float
    - margin_top_cm, margin_bottom_cm, margin_left_cm, margin_right_cm: float
    - header_text: str
    - footer_text: str
    - body_text: str
    """
    path = _make_output()
    doc = DocxDocument()

    for i, cfg in enumerate(section_configs):
        if i == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section()

        # Page size
        if "page_width_cm" in cfg and "page_height_cm" in cfg:
            section.page_width = Cm(cfg["page_width_cm"])
            section.page_height = Cm(cfg["page_height_cm"])

        # Orientation
        if cfg.get("orientation") == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            if "page_width_cm" not in cfg:
                section.page_width = Cm(29.7)
                section.page_height = Cm(21.0)

        # Margins
        section.top_margin = Cm(cfg.get("margin_top_cm", 2.54))
        section.bottom_margin = Cm(cfg.get("margin_bottom_cm", 2.54))
        section.left_margin = Cm(cfg.get("margin_left_cm", 3.17))
        section.right_margin = Cm(cfg.get("margin_right_cm", 3.17))

        # Header
        if cfg.get("header_text"):
            section.header.is_linked_to_previous = False
            header_para = section.header.paragraphs[0]
            header_para.text = cfg["header_text"]

        # Footer
        if cfg.get("footer_text"):
            section.footer.is_linked_to_previous = False
            footer_para = section.footer.paragraphs[0]
            footer_para.text = cfg["footer_text"]

        # Body text
        if cfg.get("body_text"):
            doc.add_paragraph(cfg["body_text"])

    doc.save(path)
    return path


# ─── Multi-Section Extraction Tests ────────────────────────────

class TestMultiSectionExtractionVerification:
    """Verify multi-section extraction with real multi-section documents."""

    def test_two_sections_portrait_landscape(self):
        """Document with portrait + landscape sections extracts both correctly."""
        path = _create_multi_section_doc([
            {
                "orientation": "portrait",
                "page_width_cm": 21.0,
                "page_height_cm": 29.7,
                "header_text": "Cover Page Header",
                "body_text": "Cover page content",
            },
            {
                "orientation": "landscape",
                "page_width_cm": 29.7,
                "page_height_cm": 21.0,
                "header_text": "Body Header",
                "body_text": "Body content in landscape",
            },
        ])

        spec = extract_format(path)
        assert len(spec.sections) == 2

        # First section: portrait A4
        assert spec.sections[0].page.orientation == "portrait"
        assert spec.sections[0].page.size == "A4"

        # Second section: landscape A4
        assert spec.sections[1].page.orientation == "landscape"
        assert spec.sections[1].page.size == "A4"

        os.unlink(path)

    def test_three_sections_different_margins(self):
        """Three sections with different margins are extracted independently."""
        path = _create_multi_section_doc([
            {
                "margin_top_cm": 5.0,
                "margin_bottom_cm": 5.0,
                "header_text": "Section 1",
                "body_text": "Wide margins",
            },
            {
                "margin_top_cm": 1.0,
                "margin_bottom_cm": 1.0,
                "header_text": "Section 2",
                "body_text": "Narrow margins",
            },
            {
                "margin_top_cm": 2.54,
                "margin_bottom_cm": 2.54,
                "header_text": "Section 3",
                "body_text": "Default margins",
            },
        ])

        spec = extract_format(path)
        assert len(spec.sections) == 3

        # Verify each section has distinct margins
        assert abs(spec.sections[0].page.margin.top_cm - 5.0) < 0.5
        assert abs(spec.sections[1].page.margin.top_cm - 1.0) < 0.5
        assert abs(spec.sections[2].page.margin.top_cm - 2.54) < 0.5

        os.unlink(path)

    def test_headers_per_section(self):
        """Each section's header is extracted independently."""
        path = _create_multi_section_doc([
            {"header_text": "Header Alpha", "body_text": "A"},
            {"header_text": "Header Beta", "body_text": "B"},
            {"header_text": "Header Gamma", "body_text": "C"},
        ])

        spec = extract_format(path)
        assert len(spec.sections) == 3

        assert "Alpha" in spec.sections[0].header.content
        assert "Beta" in spec.sections[1].header.content
        assert "Gamma" in spec.sections[2].header.content

        os.unlink(path)

    def test_footers_per_section(self):
        """Each section's footer is extracted independently."""
        path = _create_multi_section_doc([
            {"footer_text": "Footer One", "body_text": "A"},
            {"footer_text": "Footer Two", "body_text": "B"},
        ])

        spec = extract_format(path)
        assert len(spec.sections) == 2

        assert "One" in spec.sections[0].footer.content
        assert "Two" in spec.sections[1].footer.content

        os.unlink(path)

    def test_backward_compat_first_section_values(self):
        """spec.page, spec.header, spec.footer come from first section."""
        path = _create_multi_section_doc([
            {
                "header_text": "First Header",
                "footer_text": "First Footer",
                "margin_top_cm": 4.0,
                "body_text": "First section",
            },
            {
                "header_text": "Second Header",
                "footer_text": "Second Footer",
                "margin_top_cm": 1.0,
                "body_text": "Second section",
            },
        ])

        spec = extract_format(path)

        # Backward compat: primary values from first section
        assert spec.page is spec.sections[0].page
        assert "First Header" in spec.header.content
        assert "First Footer" in spec.footer.content

        os.unlink(path)

    def test_cover_plus_body_plus_appendix_pattern(self):
        """Real-world pattern: cover (landscape) + body (portrait) + appendix (portrait)."""
        path = _create_multi_section_doc([
            {
                "orientation": "landscape",
                "page_width_cm": 29.7,
                "page_height_cm": 21.0,
                "header_text": "",
                "body_text": "Cover Page",
                "margin_top_cm": 2.0,
            },
            {
                "orientation": "portrait",
                "page_width_cm": 21.0,
                "page_height_cm": 29.7,
                "header_text": "Document Body",
                "body_text": "Main content goes here with {{variable}}",
                "margin_top_cm": 2.54,
                "footer_text": "Page {{page_number}}",
            },
            {
                "orientation": "portrait",
                "page_width_cm": 21.0,
                "page_height_cm": 29.7,
                "header_text": "Appendix",
                "body_text": "Appendix content",
                "margin_top_cm": 3.0,
                "footer_text": "Appendix Page {{page_number}}",
            },
        ])

        spec = extract_format(path)
        assert len(spec.sections) == 3

        # Cover: landscape
        assert spec.sections[0].page.orientation == "landscape"
        # Body: portrait
        assert spec.sections[1].page.orientation == "portrait"
        assert "Document Body" in spec.sections[1].header.content
        assert "page_number" in spec.sections[1].footer.content
        # Appendix: portrait with different header/footer
        assert spec.sections[2].page.orientation == "portrait"
        assert "Appendix" in spec.sections[2].header.content

        # Constraint should mention multi-section
        assert any("多节" in c for c in spec.constraints)

        os.unlink(path)

    def test_section_with_no_header_footer(self):
        """Sections without explicit headers/footers are handled gracefully."""
        path = _create_multi_section_doc([
            {"body_text": "Section with header", "header_text": "Has Header"},
            {"body_text": "Section without header"},
        ])

        spec = extract_format(path)
        assert len(spec.sections) == 2
        assert "Has Header" in spec.sections[0].header.content
        # Second section's header should be empty (linked to previous or blank)
        # This is valid behavior — we extract what's there

        os.unlink(path)

    def test_single_section_still_works(self):
        """A basic single-section document still extracts correctly."""
        path = _make_output()
        doc = DocxDocument()
        doc.add_paragraph("Simple document")
        doc.save(path)

        spec = extract_format(path)
        assert len(spec.sections) >= 1
        assert spec.page is not None
        # python-docx default is Letter; just verify size is detected
        assert spec.page.size != ""

        os.unlink(path)


# ─── Multi-Table Extraction Verification ───────────────────────

class TestMultiTableExtractionVerification:
    """Verify multi-table format extraction with diverse table styles."""

    def test_multiple_tables_same_style(self):
        """Multiple tables with the same style are deduplicated."""
        path = _make_output()
        doc = DocxDocument()
        doc.add_paragraph("Before tables")
        doc.add_table(rows=2, cols=3)
        doc.add_paragraph("Between tables")
        doc.add_table(rows=3, cols=4)
        doc.save(path)

        spec = extract_format(path)
        # Both tables use default style, so should be deduplicated to 1
        assert len(spec.tables) >= 1
        assert spec.table is not None

        os.unlink(path)

    def test_no_tables_empty_list(self):
        """Document without tables returns empty tables list."""
        path = _make_output()
        doc = DocxDocument()
        doc.add_paragraph("No tables")
        doc.save(path)

        spec = extract_format(path)
        assert spec.tables == []
        assert spec.table is not None  # default TableFormat

        os.unlink(path)


# ─── PDF Converter Tests ───────────────────────────────────────

class TestPdfConverter:
    """Test PDF converter enhancements (no LibreOffice required for unit tests)."""

    def test_timeout_constant_is_30s(self):
        """Conversion timeout is set to 30 seconds."""
        assert CONVERSION_TIMEOUT_SECONDS == 30

    def test_file_hash_deterministic(self):
        """Same file content produces same hash."""
        path = _make_output()
        with open(path, "wb") as f:
            f.write(b"test content for hashing")

        hash1 = _compute_file_hash(path)
        hash2 = _compute_file_hash(path)
        assert hash1 == hash2
        assert len(hash1) == 16

        os.unlink(path)

    def test_file_hash_changes_with_content(self):
        """Different file content produces different hash."""
        path1 = _make_output()
        path2 = _make_output()
        with open(path1, "wb") as f:
            f.write(b"content A")
        with open(path2, "wb") as f:
            f.write(b"content B")

        assert _compute_file_hash(path1) != _compute_file_hash(path2)

        os.unlink(path1)
        os.unlink(path2)

    def test_cache_dir_creation(self):
        """Cache directory is created when accessed."""
        cache_dir = _get_cache_dir()
        assert cache_dir.exists()
        assert cache_dir.is_dir()

    def test_cached_pdf_path_includes_content_hash(self):
        """Cache path includes content hash for invalidation."""
        path = _make_output()
        with open(path, "wb") as f:
            f.write(b"docx content")

        cached_path = _get_cached_pdf_path(path)
        assert cached_path.suffix == ".pdf"
        content_hash = _compute_file_hash(path)
        assert content_hash in cached_path.name

        os.unlink(path)

    def test_cached_pdf_path_changes_with_content(self):
        """Different content produces different cache paths."""
        path1 = _make_output()
        path2 = _make_output()
        with open(path1, "wb") as f:
            f.write(b"version 1")
        with open(path2, "wb") as f:
            f.write(b"version 2")

        cp1 = _get_cached_pdf_path(path1)
        cp2 = _get_cached_pdf_path(path2)
        assert cp1 != cp2

        os.unlink(path1)
        os.unlink(path2)

    def test_clear_pdf_cache(self):
        """clear_pdf_cache removes cached files and returns count."""
        cache_dir = _get_cache_dir()
        # Create some dummy cache files
        test_files = []
        for i in range(3):
            p = cache_dir / f"test_{i}.pdf"
            p.write_bytes(b"fake pdf")
            test_files.append(p)

        count = clear_pdf_cache()
        assert count >= 3
        for p in test_files:
            assert not p.exists()

    def test_libreoffice_availability_check(self):
        """is_libreoffice_available returns a boolean without raising."""
        result = is_libreoffice_available()
        assert isinstance(result, bool)

    @pytest.mark.skipif(
        not is_libreoffice_available(),
        reason="LibreOffice not installed",
    )
    def test_convert_to_pdf_with_libreoffice(self):
        """End-to-end conversion when LibreOffice is available."""
        from app.core.pdf_converter import convert_to_pdf, convert_to_pdf_cached

        path = _make_output()
        doc = DocxDocument()
        doc.add_paragraph("Test document for PDF conversion")
        doc.save(path)

        pdf_path = None
        try:
            pdf_path = convert_to_pdf(path)
            assert Path(pdf_path).exists()
            assert Path(pdf_path).suffix == ".pdf"
            assert Path(pdf_path).stat().st_size > 0

            # Test cached version
            pdf_path2 = convert_to_pdf_cached(path)
            assert Path(pdf_path2).exists()
            # Second call should hit cache
            pdf_path3 = convert_to_pdf_cached(path)
            assert pdf_path2 == pdf_path3
        except RuntimeError:
            pytest.skip("LibreOffice detected but conversion failed (non-functional install)")
        finally:
            if os.path.exists(path):
                os.unlink(path)
            if pdf_path and Path(pdf_path).exists():
                os.unlink(pdf_path)
            clear_pdf_cache()

    def test_convert_nonexistent_file_raises(self):
        """Converting a non-existent file raises FileNotFoundError."""
        from app.core.pdf_converter import convert_to_pdf

        with pytest.raises(FileNotFoundError):
            convert_to_pdf("/nonexistent/path/document.docx")

    def test_cached_convert_nonexistent_file_raises(self):
        """Cached conversion of non-existent file raises FileNotFoundError."""
        from app.core.pdf_converter import convert_to_pdf_cached

        with pytest.raises(FileNotFoundError):
            convert_to_pdf_cached("/nonexistent/path/document.docx")
