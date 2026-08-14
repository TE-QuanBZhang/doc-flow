"""Format-preserving content replacement engine.

Replaces text/table content inside .docx by modifying OOXML nodes in place,
preserving all formatting (run properties, paragraph properties, table style,
column widths) without rebuilding the document.
"""

import re
import shutil
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

# Placeholder pattern: {{variable_name}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{\s*([^#/\s{}]+?)\s*\}\}')


class ReplacementError(Exception):
    """Raised when a replacement operation fails."""


def replace_preserving_format(
    template_path: str,
    output_path: str,
    variables: dict[str, Any],
    table_data: dict[str, list[list[Any]]] | None = None,
) -> str:
    """Replace placeholders in a document while preserving all formatting.

    Args:
        template_path: Path to the source .docx template.
        output_path: Where to save the output document.
        variables: Mapping of placeholder name -> replacement text value.
        table_data: Mapping of table placeholder name -> list of rows,
            each row is a list of cell values.

    Returns:
        Output path of the generated document.

    Raises:
        ReplacementError: If replacement fails or document is corrupted.
    """
    # Copy template to preserve original
    shutil.copy2(template_path, output_path)

    doc = DocxDocument(output_path)
    table_data = table_data or {}

    # 1. Replace text placeholders in body paragraphs
    _replace_paragraph_text(doc, variables)

    # 2. Replace placeholders inside tables + expand table data
    _replace_table_content(doc, variables, table_data)

    # 3. Save changes back to the document
    doc.save(output_path)

    # 4. Validate integrity
    _validate_document(output_path)

    return output_path


# ─── 1.2 Text replacement in paragraphs ──────────────────────

def _replace_paragraph_text(doc: DocxDocument, variables: dict[str, Any]) -> None:
    """Replace {{placeholder}} in all body paragraphs, preserving run formatting."""
    for para in doc.paragraphs:
        _replace_in_paragraph_element(para._element, variables)

    # Also process paragraphs inside tables (cell paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph_element(para._element, variables)


def _replace_in_paragraph_element(p_elem, variables: dict[str, Any]) -> None:
    """Replace placeholders within a <w:p> element, handling multi-run placeholders."""
    runs = p_elem.findall(qn('w:r'))

    # First pass: simple replacement within single runs
    for r_elem in runs:
        t_elems = r_elem.findall(qn('w:t'))
        for t_elem in t_elems:
            text = t_elem.text or ''
            if '{{' not in text:
                continue
            new_text, changed = _replace_in_text(text, variables)
            if changed:
                t_elem.text = new_text

    # Second pass: handle placeholders split across multiple runs
    _merge_split_placeholders(p_elem, variables)


def _replace_in_text(text: str, variables: dict[str, Any]) -> tuple[str, bool]:
    """Replace {{placeholders}} in a text string. Returns (new_text, changed)."""

    def replacer(match):
        var_name = match.group(1).strip()
        if var_name in variables:
            value = variables[var_name]
            return str(value)
        return match.group(0)

    new_text = PLACEHOLDER_PATTERN.sub(replacer, text)
    return new_text, new_text != text


# ─── 1.3 Cross-run placeholder merging ────────────────────────

def _merge_split_placeholders(p_elem, variables: dict[str, Any]) -> None:
    """Merge adjacent runs to resolve placeholders split across runs.

    Example: run1='{{customer', run2='_name}}' -> merge, replace, rewrite.
    Strategy: walk runs left to right; when a run ends mid-placeholder,
    append subsequent run texts until the placeholder closes.
    """
    runs = p_elem.findall(qn('w:r'))
    if len(runs) < 2:
        return

    i = 0
    while i < len(runs) - 1:
        r_elem = runs[i]
        text_i = _run_text(r_elem)

        # Check if this run contains an unclosed placeholder opener
        if '{{' in text_i and '}}' not in text_i[text_i.index('{{'):]:
            # Collect following runs until placeholder closes
            merged_text = text_i
            merged_elems = [r_elem]
            j = i + 1
            while j < len(runs) and '}}' not in merged_text:
                merged_text += _run_text(runs[j])
                merged_elems.append(runs[j])
                j += 1

            # Try replacement on merged text
            new_text, changed = _replace_in_text(merged_text, variables)
            if changed:
                # Write result back into the first run, remove the rest
                _set_run_text(r_elem, new_text)
                for extra in merged_elems[1:]:
                    p_elem.remove(extra)
                runs = p_elem.findall(qn('w:r'))
            else:
                i = j  # skip merged runs
        i += 1


def _run_text(r_elem) -> str:
    """Get full text of a run element."""
    parts = []
    for t_elem in r_elem.findall(qn('w:t')):
        parts.append(t_elem.text or '')
    return ''.join(parts)


def _set_run_text(r_elem, text: str) -> None:
    """Set a run's text, preserving the first <w:t> and clearing the rest."""
    t_elems = r_elem.findall(qn('w:t'))
    if t_elems:
        t_elems[0].text = text
        # Remove extra w:t elements
        for extra in t_elems[1:]:
            r_elem.remove(extra)


# ─── 1.4 Table content replacement ────────────────────────────

def _replace_table_content(
    doc: DocxDocument,
    variables: dict[str, Any],
    table_data: dict[str, list[list[Any]]],
) -> None:
    """Replace placeholders in tables and expand table_data into fixed rows."""
    for table in doc.tables:
        # 1. Replace placeholders in existing cell text
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph_element(para._element, variables)

        # 2. Handle table_data placeholders in header cells
        tbl_elem = table._tbl
        header_row = tbl_elem.find(qn('w:tr'))
        if header_row is None:
            continue

        # Collect header cell texts to find table placeholder name
        for cell_elem in header_row.findall(qn('w:tc')):
            cell_text = ''.join(t.text or '' for t in cell_elem.iter(qn('w:t')))
            match = PLACEHOLDER_PATTERN.search(cell_text)
            if not match:
                continue
            key = match.group(1).strip()
            if key not in table_data:
                continue

            rows = table_data[key]
            if not rows:
                continue

            # Use the second row as a template (if present), else duplicate header
            rows_elems = tbl_elem.findall(qn('w:tr'))
            if len(rows_elems) < 2:
                continue  # no template row to replicate
            template_tr = rows_elems[1]

            # Remove existing data rows (keep header + one template row)
            for tr in rows_elems[2:]:
                tbl_elem.remove(tr)

            # Fill template row with first data row, then clone for the rest
            _fill_row(template_tr, rows[0], variables)
            for data_row in rows[1:]:
                new_tr = _clone_tr(template_tr)
                _fill_row(new_tr, data_row, variables)
                tbl_elem.append(new_tr)


def _fill_row(tr_elem, row_data: list[Any], variables: dict[str, Any]) -> None:
    """Fill a table row's cells with data, replacing placeholders."""
    cells = tr_elem.findall(qn('w:tc'))
    for idx, cell_elem in enumerate(cells):
        if idx >= len(row_data):
            break
        value = str(row_data[idx])
        # Replace all w:t text in this cell with the value (first non-empty cell per paragraph)
        for t_elem in cell_elem.iter(qn('w:t')):
            t_elem.text = value
            break  # only set first text node per cell


def _clone_tr(tr_elem):
    """Clone a table row element (deep copy)."""
    import copy
    return copy.deepcopy(tr_elem)


# ─── 1.5 Integrity validation ─────────────────────────────────

def _validate_document(output_path: str) -> None:
    """Verify the output document can be opened by python-docx."""
    try:
        doc = DocxDocument(output_path)
        # Force parse of all paragraphs/tables to catch XML corruption
        _ = len(doc.paragraphs)
        for table in doc.tables:
            _ = len(table.rows)
    except Exception as e:
        raise ReplacementError(f"替换后文档损坏: {e}") from e
