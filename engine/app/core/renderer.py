"""Document renderer - unified rendering pipeline.

Primary path: format-preserving OOXML in-place replacement (true 1:1 fidelity).
Fallback 1: docxtpl (Jinja2) -- used only when preserve mode raises an exception.
Fallback 2: python-docx direct replacement -- last resort.
"""

import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument

from .template_parser import VARIABLE_PATTERN


# --- Syntax patterns ---

IF_BLOCK_START = re.compile(r'\{\{#if:(.+?)\}\}')
IF_BLOCK_ELSE = re.compile(r'\{\{else\}\}')
IF_BLOCK_END = re.compile(r'\{\{/if\}\}')
EACH_BLOCK_START = re.compile(r'\{\{#each:(.+?)\}\}')
EACH_BLOCK_END = re.compile(r'\{\{/each\}\}')
VARIABLE_NO_SPACE = re.compile(r'\{\{([^#/\s{}].*?)\}\}')

JINJA2_PATTERN = re.compile(r'\{\{.*?\}\}|\{%.*?%\}')

LOGIC_BLOCK_PATTERN = re.compile(
    r'\{%\s*(?:if|endif|else|for|endfor)\s*[^%]*%\}'
    r'|\{\{#if:[^}]+\}\}'
    r'|\{\{/if\}\}'
    r'|\{\{else\}\}'
    r'|\{\{#each:[^}]+\}\}'
    r'|\{\{/each\}\}'
)


# --- Render result ---

@dataclass
class RenderResult:
    """Result of a document generation operation."""
    output_path: str
    render_mode: str = "preserve"  # "preserve" | "docxtpl" | "python_docx"
    warnings: list[str] = field(default_factory=list)


# --- Syntax preprocessor ---

def _preprocess_to_jinja2(text: str) -> str:
    """Convert legacy Doc Flow syntax to Jinja2/docxtpl syntax."""
    text = IF_BLOCK_START.sub(r'{% if \1 %}', text)
    text = IF_BLOCK_ELSE.sub('{% else %}', text)
    text = IF_BLOCK_END.sub('{% endif %}', text)
    text = EACH_BLOCK_END.sub('{% endfor %}', text)

    def _replace_each(m):
        list_name = m.group(1).strip()
        item_var = list_name.rstrip('s') if list_name else 'item'
        if not item_var:
            item_var = 'item'
        return f'{{% for {item_var} in {list_name} %}}'

    text = EACH_BLOCK_START.sub(_replace_each, text)
    text = VARIABLE_NO_SPACE.sub(r'{{ \1 }}', text)
    return text


# --- Format-preserving render (primary path) ---

def _render_preserving_format(
    template_path: str,
    output_path: str,
    variables: dict[str, Any],
    table_data: dict[str, list[list[Any]]] | None = None,
    image_map: dict[str, Any] | None = None,
    object_map: dict[str, Any] | None = None,
) -> str:
    """Render by in-place OOXML content replacement, preserving all formatting.

    Raises on failure so the caller can fall back to another mode.
    """
    from .preserve_replace import replace_preserving_format
    from .preserve_replace_images import (
        replace_images_preserving_geometry,
        replace_embedded_objects,
    )

    replace_preserving_format(template_path, output_path, variables, table_data=table_data)

    if image_map:
        replace_images_preserving_geometry(output_path, output_path, image_map)

    if object_map:
        replace_embedded_objects(output_path, output_path, object_map)

    return output_path


# --- docxtpl render (fallback 1) ---

def _render_with_docxtpl(template_path: str, variables: dict[str, Any], output_path: str) -> bool:
    """Render using docxtpl (Jinja2). Returns True if successful."""
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return False

    try:
        doc = DocxTemplate(template_path)
        doc.render(variables)
        doc.save(output_path)
        return True
    except Exception:
        return False


# --- python-docx render (fallback 2) ---

def _render_with_python_docx(template_path: str, variables: dict[str, Any], output_path: str) -> str:
    """Fallback render using python-docx direct replacement."""
    shutil.copy2(template_path, output_path)
    doc = DocxDocument(output_path)

    for para in doc.paragraphs:
        _process_paragraph(para, variables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para, variables)

    doc.save(output_path)
    return output_path


def _process_paragraph(para, variables: dict[str, Any]) -> None:
    """Replace variables in a single paragraph."""
    if not para.text:
        return
    for run in para.runs:
        if VARIABLE_PATTERN.search(run.text):
            run.text = _replace_variables_in_text(run.text, variables)


def _replace_variables_in_text(text: str, variables: dict[str, Any]) -> str:
    """Replace {{variable}} patterns with actual values."""
    def replacer(match):
        var_name = match.group(1).strip()
        if var_name.startswith("#") or var_name.startswith("/") or var_name == "else":
            return match.group(0)
        value = _resolve_variable(var_name, variables)
        return str(value) if value is not None else match.group(0)
    return VARIABLE_PATTERN.sub(replacer, text)


def _resolve_variable(name: str, variables: dict) -> Any:
    """Resolve a variable with dot-notation support."""
    parts = name.split(".")
    value = variables
    for part in parts:
        if isinstance(value, dict):
            value = value.get(part)
        elif isinstance(value, list) and part.isdigit():
            value = value[int(part)]
        else:
            return None
    return value


# --- Logic block marker cleanup ---

def _cleanup_logic_markers(output_path: str) -> list[str]:
    """Remove residual logic block markers from document text.

    Scans all paragraphs (body + table cells) and clears text that matches
    logic block syntax (e.g. {% if %}, {{#if:cond}}, {{/each}}).

    Returns list of warnings for any markers that were cleaned.
    """
    warnings: list[str] = []
    doc = DocxDocument(output_path)
    cleaned_count = 0

    for para in doc.paragraphs:
        cleaned_count += _clean_paragraph_markers(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    cleaned_count += _clean_paragraph_markers(para)

    if cleaned_count > 0:
        warnings.append(f"cleaned {cleaned_count} residual template markers")
        doc.save(output_path)

    return warnings


def _clean_paragraph_markers(para) -> int:
    """Clean logic block markers from a single paragraph's runs. Returns count cleaned."""
    count = 0
    for run in para.runs:
        if not run.text:
            continue
        new_text, n = LOGIC_BLOCK_PATTERN.subn('', run.text)
        if n > 0:
            run.text = new_text.strip()
            count += n
    return count


# --- Main entry points ---

def generate_document(
    template_path: str,
    variables: dict[str, Any],
    output_path: str,
    use_docxtpl: bool = True,
) -> str:
    """Generate a document from a template.

    Tries format-preserving mode first, then falls back to docxtpl, then python-docx.
    Returns the output path of the generated document.
    """
    result = generate_document_with_result(template_path, variables, output_path, use_docxtpl)
    return result.output_path


def generate_document_with_result(
    template_path: str,
    variables: dict[str, Any],
    output_path: str,
    use_docxtpl: bool = True,
    table_data: dict[str, list[list[Any]]] | None = None,
    image_map: dict[str, Any] | None = None,
    object_map: dict[str, Any] | None = None,
) -> RenderResult:
    """Generate a document and return a RenderResult with mode and warnings.

    Tries format-preserving mode first. If it fails, falls back to docxtpl,
    then python-docx. The result indicates which mode was actually used.
    """
    warnings: list[str] = []

    # Primary: format-preserving mode
    try:
        _render_preserving_format(
            template_path, output_path, variables,
            table_data=table_data, image_map=image_map, object_map=object_map,
        )
        cleanup_warnings = _cleanup_logic_markers(output_path)
        warnings.extend(cleanup_warnings)
        return RenderResult(output_path=output_path, render_mode="preserve", warnings=warnings)
    except Exception as e:
        warnings.append(f"preserve mode failed, degraded: {e}")

    # Fallback 1: docxtpl
    if use_docxtpl:
        success = _render_with_docxtpl(template_path, variables, output_path)
        if success:
            cleanup_warnings = _cleanup_logic_markers(output_path)
            warnings.extend(cleanup_warnings)
            return RenderResult(output_path=output_path, render_mode="docxtpl", warnings=warnings)

    # Fallback 2: python-docx
    _render_with_python_docx(template_path, variables, output_path)
    cleanup_warnings = _cleanup_logic_markers(output_path)
    warnings.extend(cleanup_warnings)
    return RenderResult(output_path=output_path, render_mode="python_docx", warnings=warnings)


def generate_document_with_schema(
    template_path: str,
    variables: dict[str, Any],
    output_path: str,
    schema_set: Any = None,
) -> RenderResult:
    """Generate a document with schema-validated variables."""
    if schema_set:
        from .field_validator import validate_content
        result = validate_content(variables, schema_set)
        if not result.is_valid:
            errors = "; ".join(f"{e.field}: {e.reason}" for e in result.errors)
            raise ValueError(f"variable validation failed: {errors}")
        variables = result.validated_data

    return generate_document_with_result(template_path, variables, output_path)


def generate_document_preserving_format(
    template_path: str,
    output_path: str,
    variables: dict[str, Any],
    table_data: dict[str, list[list[Any]]] | None = None,
    image_map: dict[str, Any] | None = None,
    object_map: dict[str, Any] | None = None,
) -> str:
    """Generate a document by in-place content replacement, preserving all formatting.

    This is the explicit preserve-format entry point (kept for backward compatibility).
    For automatic fallback behavior, use generate_document_with_result() instead.
    """
    _render_preserving_format(
        template_path, output_path, variables,
        table_data=table_data, image_map=image_map, object_map=object_map,
    )
    _cleanup_logic_markers(output_path)
    return output_path


# --- Placeholder detection ---

def detect_unresolved_placeholders(output_path: str) -> list[str]:
    """Check generated document for unresolved placeholders (both Jinja2 and legacy)."""
    doc = DocxDocument(output_path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text_parts.append(para.text or "")

    all_text = "\n".join(text_parts)
    found = JINJA2_PATTERN.findall(all_text)
    result = []
    for v in found:
        v = v.strip()
        if v.startswith("{#"):
            continue
        # Strip outer braces/brackets to return just the inner content
        inner = v.strip("{}").strip("{%").strip("%}").strip()
        if inner and inner not in result:
            result.append(inner)
    return result
